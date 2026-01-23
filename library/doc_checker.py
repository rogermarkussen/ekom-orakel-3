"""
Kvalitetskontroll-modul for verifisering av tall i dokumenter.

Leser Word- og PDF-dokumenter og verifiserer tall mot datakildene i systemet.

Eksempel:
    from library.doc_checker import DocumentChecker, find_document

    checker = DocumentChecker()
    path = find_document()  # Finner automatisk dokumentet i data/
    findings = checker.parse_document(path)

    for finding in findings:
        matches = checker.suggest_match(finding)
        if matches:
            verified = checker.verify_number(finding, matches[0])
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from library.cache import execute_sql
from library.query_matcher import DOMAIN_SYNONYMS, extract_keywords


@dataclass
class NumberFinding:
    """Et tall funnet i dokumentet med kontekst."""
    value: float
    raw_text: str           # "91,2" eller "2,4 millioner"
    context: str            # Setning/avsnitt rundt tallet
    location: str           # "Avsnitt 5" eller "Tabell 2, rad 3"
    position: int           # For dokumentrekkefølge
    unit: Optional[str] = None  # "prosent", "millioner", etc.


@dataclass
class DataSourceMatch:
    """Foreslått datakilde-match for et tall."""
    source: str             # "ekom", "dekning_tek", etc.
    sql: str                # SQL for verifisering
    expected_value: float
    confidence: float       # 0-1
    filters: dict           # For visning til bruker
    description: str = ""   # Lesbar beskrivelse av spørringen


@dataclass
class LearnedPattern:
    """Lært mønster fra bruker-interaksjon."""
    context_keywords: list[str]
    source: str
    filters: dict
    sql_template: str = ""


# Heuristikker for auto-matching basert på kontekst
CONTEXT_HEURISTICS = {
    # Teknologidekning
    "fiber": {
        "source": "dekning_tek",
        "filters": {"tek": "fiber"},
        "unit": "prosent",
    },
    "fiberdekning": {
        "source": "dekning_tek",
        "filters": {"tek": "fiber"},
        "unit": "prosent",
    },
    "5g": {
        "source": "dekning_tek",
        "filters": {"tek": "5g"},
        "unit": "prosent",
    },
    "5g-dekning": {
        "source": "dekning_tek",
        "filters": {"tek": "5g"},
        "unit": "prosent",
    },
    "4g": {
        "source": "dekning_tek",
        "filters": {"tek": "4g"},
        "unit": "prosent",
    },
    "kabel": {
        "source": "dekning_tek",
        "filters": {"tek": "kabel"},
        "unit": "prosent",
    },
    "ftb": {
        "source": "dekning_tek",
        "filters": {"tek": "ftb"},
        "unit": "prosent",
    },
    "fast trådløst": {
        "source": "dekning_tek",
        "filters": {"tek": "ftb"},
        "unit": "prosent",
    },

    # Ekom - Abonnement
    "abonnement": {
        "source": "ekom",
        "filters": {"hg": "Abonnement"},
        "unit": "antall",
    },
    "abonnenter": {
        "source": "ekom",
        "filters": {"hg": "Abonnement"},
        "unit": "antall",
    },
    "mobilabonnement": {
        "source": "ekom",
        "filters": {"hk": "Mobiltjenester", "hg": "Abonnement"},
        "unit": "antall",
    },
    "bredbåndsabonnement": {
        "source": "ekom",
        "filters": {"hk": "Fast bredbånd", "hg": "Abonnement"},
        "unit": "antall",
    },

    # Ekom - Inntekter
    "inntekt": {
        "source": "ekom",
        "filters": {"hg": "Inntekter"},
        "unit": "kroner",
    },
    "omsetning": {
        "source": "ekom",
        "filters": {"hg": "Inntekter"},
        "unit": "kroner",
    },
    "inntekter": {
        "source": "ekom",
        "filters": {"hg": "Inntekter"},
        "unit": "kroner",
    },
    "investeringer": {
        "source": "ekom",
        "filters": {"hk": "Investeringer"},
        "unit": "kroner",
    },

    # Ekom - Markedsandeler
    "markedsandel": {
        "source": "ekom",
        "filters": {},
        "unit": "prosent",
    },
    "markedsandeler": {
        "source": "ekom",
        "filters": {},
        "unit": "prosent",
    },
    "tilbyder": {
        "source": "ekom",
        "filters": {},
        "unit": None,
    },
    "telenor": {
        "source": "ekom",
        "filters": {"fusnavn_contains": "Telenor"},
        "unit": None,
    },
    "telia": {
        "source": "ekom",
        "filters": {"fusnavn_contains": "Telia"},
        "unit": None,
    },

    # Husstander
    "husstander": {
        "source": "adr",
        "filters": {},
        "unit": "antall",
    },
    "husholdninger": {
        "source": "adr",
        "filters": {},
        "unit": "antall",
    },

    # Geografiske filtre (modifiserer andre heuristikker)
    "spredtbygd": {
        "source": None,  # Modifikator
        "filters": {"geo": "spredtbygd"},
        "unit": None,
    },
    "tettbygd": {
        "source": None,
        "filters": {"geo": "tettbygd"},
        "unit": None,
    },
    "tettsted": {
        "source": None,
        "filters": {"geo": "tettbygd"},
        "unit": None,
    },
}

# Støttede filtyper
SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


def find_document(data_dir: Path | str = "data") -> Path:
    """
    Finn dokumentet i data-mappen.

    Forventer at det kun er ett dokument (Word eller PDF) i mappen.

    Args:
        data_dir: Sti til data-mappen (default: "data")

    Returns:
        Path til dokumentet

    Raises:
        FileNotFoundError: Hvis ingen dokument finnes
        ValueError: Hvis flere dokumenter finnes
    """
    data_path = Path(data_dir)

    if not data_path.exists():
        raise FileNotFoundError(f"Data-mappen '{data_dir}' finnes ikke")

    documents = [
        f for f in data_path.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if len(documents) == 0:
        raise FileNotFoundError(
            f"Ingen dokumenter funnet i '{data_dir}'. "
            f"Støttede formater: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if len(documents) > 1:
        doc_names = [d.name for d in documents]
        raise ValueError(
            f"Flere dokumenter funnet i '{data_dir}': {doc_names}. "
            "Det skal kun være ett dokument."
        )

    return documents[0]


# Multipler for enheter
UNIT_MULTIPLIERS = {
    "millioner": 1_000_000,
    "mill": 1_000_000,
    "mill.": 1_000_000,
    "million": 1_000_000,
    "milliarder": 1_000_000_000,
    "mrd": 1_000_000_000,
    "mrd.": 1_000_000_000,
    "milliard": 1_000_000_000,
    "tusen": 1_000,
}


def parse_norwegian_number(text: str) -> tuple[float, Optional[str]]:
    """
    Parser norsk tallformat til float.

    Eksempler:
        "91,2" -> (91.2, None)
        "2,4 millioner" -> (2400000.0, "millioner")
        "1 234 567" -> (1234567.0, None)

    Returnerer (value, unit).
    """
    text = text.strip()

    # Finn enhet først - sjekk lengste enheter først for å unngå delmatch
    unit = None
    multiplier = 1.0
    text_lower = text.lower()
    for unit_name, mult in sorted(UNIT_MULTIPLIERS.items(), key=lambda x: len(x[0]), reverse=True):
        if unit_name in text_lower:
            unit = unit_name
            multiplier = mult
            # Fjern enheten fra teksten
            text = re.sub(rf'\s*{re.escape(unit_name)}\.?\s*', '', text, flags=re.IGNORECASE)
            break

    # Fjern prosent-tegn
    if "%" in text or "prosent" in text.lower():
        text = text.replace("%", "").replace("prosent", "").strip()
        if unit is None:
            unit = "prosent"

    # Håndter norsk tallformat
    # Fjern mellomrom (tusenskilletegn)
    text = text.replace(" ", "").replace("\u00a0", "")

    # Konverter komma til punktum (desimaltegn)
    # Men pass på at vi ikke har punktum som tusenskilletegn
    if "," in text and "." not in text:
        text = text.replace(",", ".")
    elif "." in text and "," in text:
        # Begge finnes - punktum er tusenskilletegn, komma er desimal
        text = text.replace(".", "").replace(",", ".")

    try:
        value = float(text) * multiplier
        return value, unit
    except ValueError:
        return 0.0, None


class DocumentChecker:
    """Hovedklasse for dokumentkontroll."""

    def __init__(self):
        self.learned_patterns: list[LearnedPattern] = []

    def parse_document(self, path: Path) -> list[NumberFinding]:
        """
        Ekstraher alle tall med kontekst fra et dokument (Word eller PDF).

        Args:
            path: Sti til .docx eller .pdf fil

        Returns:
            Liste med NumberFinding objekter
        """
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return self._parse_word(path)
        elif suffix == ".pdf":
            return self._parse_pdf(path)
        else:
            raise ValueError(
                f"Ikke støttet filtype: {suffix}. "
                f"Støttede formater: {', '.join(SUPPORTED_EXTENSIONS)}"
            )

    def _parse_word(self, path: Path) -> list[NumberFinding]:
        """Parser Word-dokument (.docx)."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx er ikke installert. Kjør: uv add python-docx"
            )

        doc = Document(path)
        findings: list[NumberFinding] = []
        position = 0

        # Parser avsnitt
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # Finn tall i avsnittet
            para_findings = self._extract_numbers_from_text(
                text,
                location=f"Avsnitt {para_idx + 1}",
                base_position=position
            )
            findings.extend(para_findings)
            position += len(para_findings)

        # Parser tabeller
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    cell_findings = self._extract_numbers_from_text(
                        text,
                        location=f"Tabell {table_idx + 1}, rad {row_idx + 1}, celle {cell_idx + 1}",
                        base_position=position
                    )
                    findings.extend(cell_findings)
                    position += len(cell_findings)

        return findings

    def _parse_pdf(self, path: Path) -> list[NumberFinding]:
        """Parser PDF-dokument."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber er ikke installert. Kjør: uv add pdfplumber"
            )

        findings: list[NumberFinding] = []
        position = 0

        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Ekstraher tekst fra siden
                text = page.extract_text() or ""

                # Del opp i avsnitt (dobble linjeskift)
                paragraphs = text.split("\n\n")

                for para_idx, paragraph in enumerate(paragraphs):
                    paragraph = paragraph.strip()
                    if not paragraph:
                        continue

                    # Finn tall i avsnittet
                    para_findings = self._extract_numbers_from_text(
                        paragraph,
                        location=f"Side {page_num}, avsnitt {para_idx + 1}",
                        base_position=position
                    )
                    findings.extend(para_findings)
                    position += len(para_findings)

                # Ekstraher også tabeller
                tables = page.extract_tables() or []
                for table_idx, table in enumerate(tables):
                    for row_idx, row in enumerate(table):
                        for cell_idx, cell in enumerate(row or []):
                            if not cell:
                                continue
                            cell = str(cell).strip()
                            if not cell:
                                continue

                            cell_findings = self._extract_numbers_from_text(
                                cell,
                                location=f"Side {page_num}, tabell {table_idx + 1}, rad {row_idx + 1}",
                                base_position=position
                            )
                            findings.extend(cell_findings)
                            position += len(cell_findings)

        return findings

    def _extract_numbers_from_text(
        self,
        text: str,
        location: str,
        base_position: int
    ) -> list[NumberFinding]:
        """Ekstraher tall fra en tekststreng."""
        findings = []

        # Regex for å finne tall med valgfri enhet
        # Matcher: 91,2 | 2,4 millioner | 1 234 567 | 85% | 2.4 mill.
        # Viktig: Lengre matcher (milliarder) må komme før kortere (mill)
        number_pattern = r"""
            (\d[\d\s]*(?:[,.]\d+)?)  # Tall med mulige mellomrom og desimaler
            \s*
            (%|prosent|milliarder|milliard|mrd\.?|millioner|million|mill\.?|tusen)?  # Valgfri enhet
        """

        for match in re.finditer(number_pattern, text, re.VERBOSE | re.IGNORECASE):
            raw_text = match.group(0).strip()
            has_unit = match.group(2) is not None

            # Ignorer veldig korte tall (sannsynligvis ikke statistikk)
            if len(raw_text) < 2:
                continue

            # Ignorer årstall (2020, 2024, etc.)
            if re.match(r'^20\d{2}$', raw_text.replace(" ", "")):
                continue

            value, unit = parse_norwegian_number(raw_text)

            if value == 0:
                continue

            # Ignorer seksjons-/kapittel-numre (1, 2, 3, 4.1, 4.2, etc.)
            if re.match(r'^\d+(\.\d+)?$', raw_text.replace(" ", "")) and value < 50 and not has_unit:
                # Sjekk om det ser ut som en seksjonsnummer (tall i starten av tekst)
                text_stripped = text.strip()
                if text_stripped.startswith(raw_text):
                    continue
                # Også ignorer enkelt-siffer tall uten kontekst
                if value < 10 and not has_unit:
                    continue

            # Ignorer sidetall (enkeltstående tall mellom 1-300 uten enhet)
            if value <= 300 and value == int(value) and not has_unit:
                # Sjekk om hele teksten bare er et tall (typisk sidetall)
                if text.strip() == raw_text:
                    continue
                # Sjekk om tallet er alene med kun whitespace/tab rundt
                if re.match(r'^\s*\d+\s*$', text):
                    continue

            # Finn kontekst (hele setningen)
            # Finn setningen tallet er i
            sentences = re.split(r'[.!?]', text)
            context = text  # Fallback til hele teksten
            for sentence in sentences:
                if raw_text in sentence:
                    context = sentence.strip()
                    break

            # Ignorer tall hvor konteksten hovedsakelig er en innholdsfortegnelse-linje
            # (typisk: "Kapittel X ... sidetall" eller "Figur 5 Beskrivelse\tsidetall")
            if re.match(r'^[\d\.]+\s+[A-ZÆØÅa-zæøå\s-]+\t+\d+$', context):
                continue

            # Ignorer sidetall i innholdsfortegnelse (tekst + tab + tall)
            if re.search(r'\t\d+$', context) and not has_unit:
                # Tallet er sannsynligvis sidetallet på slutten
                if context.endswith(raw_text):
                    continue

            # Ignorer figur-/tabell-numre (Figur 5, Tabell 3, etc.)
            if re.match(r'^(Figur|Tabell|Diagram|Graf)\s+\d+', context, re.IGNORECASE):
                # Sjekk om dette tallet er figurnummeret
                fig_match = re.match(r'^(Figur|Tabell|Diagram|Graf)\s+(\d+)', context, re.IGNORECASE)
                if fig_match and fig_match.group(2) == raw_text:
                    continue

            # Ignorer tall under 100 uten enhet når konteksten er kort (typisk indeks/ref)
            if value < 100 and not has_unit and len(context) < 50:
                continue

            findings.append(NumberFinding(
                value=value,
                raw_text=raw_text,
                context=context,
                location=location,
                position=base_position + len(findings),
                unit=unit
            ))

        return findings

    def suggest_match(self, finding: NumberFinding) -> list[DataSourceMatch]:
        """
        Foreslå datakilde-matcher for et tall basert på kontekst.

        Args:
            finding: Tallet med kontekst

        Returns:
            Liste med mulige matcher, sortert etter confidence
        """
        matches = []
        context_lower = finding.context.lower()

        # Ekstraher nøkkelord fra kontekst
        keywords = extract_keywords(finding.context)

        # Sjekk lærte mønstre først
        for pattern in self.learned_patterns:
            if all(kw.lower() in context_lower for kw in pattern.context_keywords):
                match = self._create_match_from_pattern(finding, pattern)
                if match:
                    match.confidence = 0.95  # Høy confidence for lærte mønstre
                    matches.append(match)

        # Spesifikke mønstre har høyere prioritet enn generelle
        # Sjekk kombinerte begreper først
        # Format: (terms, match_mode, source, filters)
        # match_mode: "any" = et av termene matcher, "all" = alle termene må matche
        combined_patterns = [
            # Fiber + omsetning/inntekt = ekom Inntekter (sjekkes FØR abonnement)
            (["fiber", "omsetning"], "all", "ekom",
             {"hk": "Fast bredbånd", "hg": "Inntekter", "tek": "Fiber"}),
            (["fiber", "inntekt"], "all", "ekom",
             {"hk": "Fast bredbånd", "hg": "Inntekter", "tek": "Fiber"}),
            (["fiber", "kroner"], "all", "ekom",
             {"hk": "Fast bredbånd", "hg": "Inntekter", "tek": "Fiber"}),
            (["fiber", "milliard"], "all", "ekom",
             {"hk": "Fast bredbånd", "hg": "Inntekter", "tek": "Fiber"}),
            # Fiber + abonnement = ekom med tek=Fiber
            (["fiberabonnement", "fiber-abonnement"], "any", "ekom",
             {"hk": "Fast bredbånd", "hg": "Abonnement", "tek": "Fiber"}),
            (["fiber", "abonnement"], "all", "ekom",
             {"hk": "Fast bredbånd", "hg": "Abonnement", "tek": "Fiber"}),
            # Mobil-spesifikke
            (["mobilabonnement", "mobil-abonnement"], "any", "ekom",
             {"hk": "Mobiltjenester", "hg": "Abonnement"}),
            (["mobiltelefoni", "abonnement"], "all", "ekom",
             {"hk": "Mobiltjenester", "dk": "Mobiltelefoni", "hg": "Abonnement"}),
            # TV-spesifikke
            (["tv-abonnement", "tv abonnement"], "any", "ekom",
             {"hk": "TV-tjenester", "hg": "Abonnement"}),
            # Dekning (kun hvis ikke abonnement/inntekt)
            (["fiberdekning", "fiber-dekning"], "any", "dekning_tek",
             {"tek": "fiber"}),
            (["fiber", "dekning"], "all", "dekning_tek",
             {"tek": "fiber"}),
        ]

        for terms, match_mode, source, filters in combined_patterns:
            # Sjekk om termene matcher
            if match_mode == "any":
                term_match = any(t in context_lower for t in terms)
            else:  # "all"
                term_match = all(t in context_lower for t in terms)

            if term_match:
                # Sjekk at dette ikke er en dekning-kilde når det handler om abonnement
                if source == "dekning_tek" and any(x in context_lower for x in ["abonnement", "omsetning", "inntekt"]):
                    continue

                # Hopp over ekom-match for prosent-verdier som ikke er markedsandeler
                if source == "ekom" and finding.unit == "prosent":
                    if "markedsandel" not in context_lower and "andel" not in context_lower:
                        # Prosent-verdi uten markedsandel-kontekst - bruk dekning istedet
                        continue

                match = self._build_match(finding, source, filters, keywords)
                if match:
                    matches.append(match)
                    break  # Bruk første match fra kombinerte mønstre

        # Hvis ingen kombinerte match, sjekk enkle heuristikker
        if not matches:
            detected_source = None
            detected_filters: dict = {}

            # Prosent-verdier som ikke er markedsandeler -> prøv dekning
            if finding.unit == "prosent" and "markedsandel" not in context_lower and "andel" not in context_lower:
                # Prøv dekning_tek basert på teknologi i konteksten
                for tek_name in ["fiber", "5g", "4g", "kabel", "ftb"]:
                    if tek_name in context_lower:
                        detected_source = "dekning_tek"
                        detected_filters = {"tek": tek_name}
                        break

            # Prioriter ekom-kilder for abonnement/inntekt (unntatt prosent)
            elif any(x in context_lower for x in ["abonnement", "abonnenter", "omsetning", "inntekt", "kroner", "milliarder"]):
                detected_source = "ekom"
                if "abonnement" in context_lower or "abonnenter" in context_lower:
                    detected_filters["hg"] = "Abonnement"
                elif any(x in context_lower for x in ["omsetning", "inntekt", "kroner", "milliarder"]):
                    detected_filters["hg"] = "Inntekter"

                # Legg til teknologi hvis nevnt
                if "fiber" in context_lower:
                    detected_filters["tek"] = "Fiber"
                    detected_filters["hk"] = "Fast bredbånd"
                elif "mobil" in context_lower:
                    detected_filters["hk"] = "Mobiltjenester"

            # Fallback til standard heuristikker
            if not detected_source:
                for keyword, heuristic in CONTEXT_HEURISTICS.items():
                    if keyword.lower() in context_lower:
                        if heuristic["source"]:
                            detected_source = heuristic["source"]
                        detected_filters.update(heuristic["filters"])

            if detected_source:
                match = self._build_match(
                    finding, detected_source, detected_filters, keywords
                )
                if match:
                    matches.append(match)

        # Sorter etter confidence
        matches.sort(key=lambda m: m.confidence, reverse=True)

        return matches

    def _build_match(
        self,
        finding: NumberFinding,
        source: str,
        filters: dict,
        keywords: list[str]
    ) -> Optional[DataSourceMatch]:
        """Bygg en DataSourceMatch med SQL og forventet verdi."""

        # Bestem år - default 2024
        year = 2024

        if source == "dekning_tek":
            return self._build_dekning_tek_match(finding, filters, year)
        elif source == "dekning_hast":
            return self._build_dekning_hast_match(finding, filters, year)
        elif source == "ekom":
            return self._build_ekom_match(finding, filters, year)
        elif source == "adr":
            return self._build_adr_match(finding, filters, year)

        return None

    def _build_dekning_tek_match(
        self,
        finding: NumberFinding,
        filters: dict,
        year: int
    ) -> Optional[DataSourceMatch]:
        """Bygg match for dekning_tek."""
        tek = filters.get("tek", "fiber")
        geo = filters.get("geo", "totalt")

        sql = f"""
SELECT ROUND(dekning * 100, 1) as prosent
FROM dekning_tek
WHERE tek = '{tek}'
  AND fylke = 'NASJONALT'
  AND geo = '{geo}'
  AND ar = {year}
"""

        try:
            result = execute_sql(sql)
            if result.height > 0:
                expected = result.item(0, 0)

                # Beregn confidence basert på verdidifferanse
                if finding.value > 1:  # Prosent-verdi
                    diff = abs(finding.value - expected)
                else:  # Desimalverdi (0-1)
                    diff = abs(finding.value * 100 - expected)

                if diff < 1:
                    confidence = 0.9
                elif diff < 5:
                    confidence = 0.7
                elif diff < 10:
                    confidence = 0.5
                else:
                    confidence = 0.3

                return DataSourceMatch(
                    source="dekning_tek",
                    sql=sql.strip(),
                    expected_value=expected,
                    confidence=confidence,
                    filters={"tek": tek, "geo": geo, "ar": year},
                    description=f"{tek}-dekning {geo} {year}"
                )
        except Exception:
            pass

        return None

    def _build_dekning_hast_match(
        self,
        finding: NumberFinding,
        filters: dict,
        year: int
    ) -> Optional[DataSourceMatch]:
        """Bygg match for dekning_hast."""
        ned = filters.get("ned", 100)
        opp = filters.get("opp", 100)
        geo = filters.get("geo", "totalt")

        sql = f"""
SELECT ROUND(dekning * 100, 1) as prosent
FROM dekning_hast
WHERE fylke = 'NASJONALT'
  AND geo = '{geo}'
  AND ned = {ned}
  AND opp = {opp}
  AND ar = {year}
"""

        try:
            result = execute_sql(sql)
            if result.height > 0:
                expected = result.item(0, 0)
                diff = abs(finding.value - expected) if finding.value > 1 else abs(finding.value * 100 - expected)
                confidence = 0.9 if diff < 1 else 0.7 if diff < 5 else 0.5

                return DataSourceMatch(
                    source="dekning_hast",
                    sql=sql.strip(),
                    expected_value=expected,
                    confidence=confidence,
                    filters={"ned": ned, "opp": opp, "geo": geo, "ar": year},
                    description=f"{ned}/{opp} Mbit dekning {geo} {year}"
                )
        except Exception:
            pass

        return None

    def _build_ekom_match(
        self,
        finding: NumberFinding,
        filters: dict,
        year: int
    ) -> Optional[DataSourceMatch]:
        """Bygg match for ekom."""
        hk = filters.get("hk", "Fast bredbånd")
        hg = filters.get("hg", "Abonnement")
        rapport = f"{year}-Helår"

        where_clauses = [
            f"hk = '{hk}'",
            f"hg = '{hg}'",
            "tp = 'Sum'",
            "sk IN ('Sluttbruker', 'Ingen')",
            f"rapport = '{rapport}'"
        ]

        if "tek" in filters:
            where_clauses.append(f"tek = '{filters['tek']}'")
        if "ms" in filters:
            where_clauses.append(f"ms = '{filters['ms']}'")

        where_sql = " AND ".join(where_clauses)

        # Bestem divisor basert på forventet størrelsesorden
        if finding.value > 1_000_000:
            divisor = 1
            unit_label = ""
        elif finding.value > 1000:
            divisor = 1000
            unit_label = " (tusen)"
        else:
            divisor = 1_000_000
            unit_label = " (millioner)"

        sql = f"""
SELECT ROUND(SUM(svar) / {divisor}, 1) as verdi
FROM ekom
WHERE {where_sql}
"""

        try:
            result = execute_sql(sql)
            if result.height > 0:
                expected = result.item(0, 0)

                # Juster finding.value hvis den er i millioner
                compare_value = finding.value
                if finding.unit in ("millioner", "mill", "mill."):
                    compare_value = finding.value / 1_000_000
                elif finding.unit in ("milliarder", "mrd", "mrd."):
                    compare_value = finding.value / 1_000_000_000

                diff_pct = abs(compare_value - expected) / max(expected, 1) * 100
                confidence = 0.9 if diff_pct < 1 else 0.7 if diff_pct < 5 else 0.5

                return DataSourceMatch(
                    source="ekom",
                    sql=sql.strip(),
                    expected_value=expected,
                    confidence=confidence,
                    filters=filters | {"rapport": rapport},
                    description=f"{hk} - {hg}{unit_label} {rapport}"
                )
        except Exception:
            pass

        return None

    def _build_adr_match(
        self,
        finding: NumberFinding,
        filters: dict,
        year: int
    ) -> Optional[DataSourceMatch]:
        """Bygg match for adr (husstander)."""
        geo_filter = ""
        if "geo" in filters:
            if filters["geo"] == "spredtbygd":
                geo_filter = " WHERE NOT ertett"
            elif filters["geo"] == "tettbygd":
                geo_filter = " WHERE ertett"

        sql = f"""
SELECT ROUND(SUM(hus) / 1000000.0, 2) as mill_hus
FROM adr_{year}
{geo_filter}
"""

        try:
            result = execute_sql(sql)
            if result.height > 0:
                expected = result.item(0, 0)

                compare_value = finding.value
                if finding.unit in ("millioner", "mill", "mill."):
                    compare_value = finding.value / 1_000_000

                diff = abs(compare_value - expected)
                confidence = 0.9 if diff < 0.1 else 0.7 if diff < 0.3 else 0.5

                return DataSourceMatch(
                    source="adr",
                    sql=sql.strip(),
                    expected_value=expected,
                    confidence=confidence,
                    filters={"year": year} | filters,
                    description=f"Husstander {year}"
                )
        except Exception:
            pass

        return None

    def _create_match_from_pattern(
        self,
        finding: NumberFinding,
        pattern: LearnedPattern
    ) -> Optional[DataSourceMatch]:
        """Lag match fra et lært mønster."""
        return self._build_match(
            finding,
            pattern.source,
            pattern.filters,
            pattern.context_keywords
        )

    def verify_number(
        self,
        finding: NumberFinding,
        match: DataSourceMatch,
        tolerance: float = 0.5
    ) -> tuple[bool, float]:
        """
        Verifiser et tall mot forventet verdi.

        Args:
            finding: Tallet fra dokumentet
            match: Matchen å verifisere mot
            tolerance: Akseptert avvik i prosent

        Returns:
            (verified, diff_percent)
        """
        expected = match.expected_value
        actual = finding.value

        # Juster for enhet
        if finding.unit == "prosent" and expected <= 100:
            pass  # Begge er prosent
        elif finding.unit in ("millioner", "mill", "mill."):
            actual = finding.value / 1_000_000
        elif finding.unit in ("milliarder", "mrd", "mrd."):
            actual = finding.value / 1_000_000_000

        if expected == 0:
            return actual == 0, 0.0

        diff_percent = abs(actual - expected) / expected * 100
        verified = diff_percent <= tolerance

        return verified, diff_percent

    def learn_pattern(
        self,
        context: str,
        source: str,
        filters: dict,
        sql_template: str = ""
    ):
        """
        Lær et nytt mønster fra bruker-interaksjon.

        Args:
            context: Kontekst-teksten
            source: Datakilde (dekning_tek, ekom, etc.)
            filters: Filtre brukt
            sql_template: Valgfri SQL-mal
        """
        keywords = extract_keywords(context)

        pattern = LearnedPattern(
            context_keywords=keywords,
            source=source,
            filters=filters,
            sql_template=sql_template
        )

        self.learned_patterns.append(pattern)

    def get_verification_summary(
        self,
        findings: list[NumberFinding],
        results: list[tuple[NumberFinding, Optional[DataSourceMatch], bool]]
    ) -> dict:
        """
        Generer oppsummering av verifisering.

        Args:
            findings: Alle tall funnet
            results: Liste med (finding, match, verified) tupler

        Returns:
            Oppsummeringsdict
        """
        verified = sum(1 for _, _, v in results if v)
        mismatched = sum(1 for _, m, v in results if m and not v)
        skipped = sum(1 for _, m, _ in results if m is None)
        no_match = len(findings) - len(results)

        return {
            "totalt": len(findings),
            "verifisert": verified,
            "avvik": mismatched,
            "hoppet_over": skipped,
            "ingen_match": no_match,
            "prosent_verifisert": round(verified / len(findings) * 100, 1) if findings else 0
        }
